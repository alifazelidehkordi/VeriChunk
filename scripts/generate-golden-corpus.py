#!/usr/bin/env python3
"""Generate deterministic phase-zero golden fixtures.

The corpus intentionally describes desired behavior, including known gaps in the
current implementation. Run from anywhere; output is written under tests/golden.
"""

from __future__ import annotations

import json
from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from doc_splitter.ir.models import DocumentIR, DocumentMeta, Element  # noqa: E402
from doc_splitter.ir.serialize import save_ir  # noqa: E402

GOLDEN = ROOT / "tests" / "golden"
IR_DIR = GOLDEN / "ir"
SOURCE_DIR = GOLDEN / "source"


def _save(name: str, elements: list[Element]) -> str:
    ir = DocumentIR(elements=elements, meta=DocumentMeta(source_file=f"{name}.pdf"))
    ir.recompute_word_counts()
    ir.meta.estimated_total_pages = max(
        (element.page_number or 1) for element in elements
    )
    relative = Path("ir") / f"{name}.json"
    save_ir(ir, GOLDEN / relative)
    return relative.as_posix()


def _p(element_id: str, page: int, text: str) -> Element:
    return Element(id=element_id, type="paragraph", page_number=page, text=text)


def _h(element_id: str, page: int, text: str, level: int = 1) -> Element:
    return Element(
        id=element_id,
        type="heading",
        page_number=page,
        text=text,
        level=level,
    )


def _build_ir_cases() -> list[dict]:
    cases: list[dict] = []

    source = _save(
        "topic_change_with_heading",
        [
            _h("el-001", 1, "CELL BIOLOGY"),
            _p("el-002", 1, "Cells use membranes to regulate transport."),
            _p("el-003", 2, "Membrane proteins support diffusion and signaling."),
            _p("el-004", 3, "Organelles divide metabolic responsibilities."),
            _p("el-005", 4, "This concludes the cell biology learning unit."),
            _h("el-006", 5, "RENAL PHYSIOLOGY"),
            _p("el-007", 5, "The nephron filters plasma and adjusts electrolytes."),
            _p("el-008", 6, "Tubular transport determines final urine composition."),
        ],
    )
    cases.append(
        {
            "id": "topic-change-with-heading",
            "kind": "ir",
            "description": "A clear independent topic begins at a major heading.",
            "source": source,
            "expected": {"topic_boundaries_after": ["el-005"]},
        }
    )

    source = _save(
        "topic_change_without_heading",
        [
            _p("el-001", 1, "Glycolysis converts glucose into pyruvate in the cytosol."),
            _p("el-002", 2, "ATP investment primes glucose for later energy capture."),
            _p("el-003", 3, "NADH carries reducing power from the oxidation step."),
            _p("el-004", 4, "The pathway ends with net ATP and pyruvate production."),
            _p("el-005", 5, "A valid contract requires offer, acceptance, and consideration."),
            _p("el-006", 6, "Capacity and legality also affect enforceability."),
            _p("el-007", 7, "A material breach can excuse further performance."),
            _p("el-008", 8, "Damages generally aim to protect expectation interests."),
        ],
    )
    cases.append(
        {
            "id": "topic-change-without-heading",
            "kind": "ir",
            "description": "The learning objective changes without any heading marker.",
            "source": source,
            "expected": {"topic_boundaries_after": ["el-004"]},
        }
    )

    source = _save(
        "subheading_same_topic",
        [
            _h("el-001", 1, "ACID BASE BALANCE"),
            _p("el-002", 1, "Blood pH depends on respiratory and renal regulation."),
            _p("el-003", 2, "The bicarbonate buffer system limits acute pH shifts."),
            _p("el-004", 3, "Compensation changes carbon dioxide or bicarbonate."),
            _h("el-005", 4, "Worked Example", level=2),
            _p("el-006", 4, "Interpret pH, carbon dioxide, and bicarbonate in order."),
            _p("el-007", 5, "The example applies the same acid base framework."),
        ],
    )
    cases.append(
        {
            "id": "subheading-same-topic",
            "kind": "ir",
            "description": "A worked-example subheading must not split one coherent topic.",
            "source": source,
            "expected": {
                "topic_boundaries_after": [],
                "must_not_split_after": ["el-004"],
            },
        }
    )

    source = _save(
        "early_topic_change_before_minimum",
        [
            _h("el-001", 1, "INTRODUCTION TO STATISTICS"),
            _p("el-002", 1, "Descriptive statistics summarize observed samples."),
            _p("el-003", 2, "Means and medians describe central tendency."),
            _h("el-004", 3, "CARDIAC ELECTROPHYSIOLOGY"),
            _p("el-005", 3, "Ion channels generate the cardiac action potential."),
            _p("el-006", 4, "Conduction coordinates atrial and ventricular activation."),
        ],
    )
    cases.append(
        {
            "id": "early-topic-change-before-minimum",
            "kind": "ir",
            "description": "A real topic change at page 3 overrides a five-page target minimum.",
            "source": source,
            "config": {"target_min_pages": 5},
            "expected": {"topic_boundaries_after": ["el-003"]},
        }
    )

    continuous_17 = [_h("el-001", 1, "IMMUNE RESPONSE")]
    for page in range(1, 18):
        continuous_17.append(
            _p(
                f"el-{page + 1:03d}",
                page,
                f"Page {page}: the same immune-response explanation continues with connected evidence.",
            )
        )
    source = _save("continuous_topic_17_pages", continuous_17)
    cases.append(
        {
            "id": "continuous-topic-17-pages",
            "kind": "ir",
            "description": "One topic may exceed 13 pages only when semantic continuity is confirmed.",
            "source": source,
            "expected": {
                "topic_boundaries_after": [],
                "page_policy": {"allowed_single_chunk_pages": 17, "requires_extension_evidence": True},
            },
        }
    )

    continuous_25 = [_h("el-001", 1, "LONG CONTINUOUS DERIVATION")]
    for page in range(1, 26):
        continuous_25.append(
            _p(
                f"el-{page + 1:03d}",
                page,
                f"Page {page}: the same derivation continues without introducing a new topic.",
            )
        )
    source = _save("continuous_topic_25_pages", continuous_25)
    cases.append(
        {
            "id": "continuous-topic-25-pages",
            "kind": "ir",
            "description": "A continuous topic still requires a forced continuation split at 20 pages.",
            "source": source,
            "expected": {
                "topic_boundaries_after": [],
                "page_policy": {
                    "hard_max_pages": 20,
                    "forced_size_split": True,
                    "continues_to_next": True,
                },
            },
        }
    )

    source = _save(
        "table_on_topic_boundary",
        [
            _h("el-001", 1, "ANTIBIOTIC CLASSES"),
            _p("el-002", 1, "Antibiotics can be grouped by target and mechanism."),
            Element(
                id="el-003",
                type="table",
                page_number=3,
                rows=[
                    ["Class", "Target"],
                    ["Beta-lactam", "Cell wall"],
                    ["Macrolide", "Ribosome"],
                ],
            ),
            _h("el-004", 4, "VIRAL REPLICATION"),
            _p("el-005", 4, "Viruses depend on host machinery to reproduce."),
        ],
    )
    cases.append(
        {
            "id": "table-on-topic-boundary",
            "kind": "ir",
            "description": "A completed table stays with the topic before the transition.",
            "source": source,
            "expected": {"topic_boundaries_after": ["el-003"]},
        }
    )

    source = _save(
        "image_on_topic_boundary",
        [
            _h("el-001", 1, "NEURAL PATHWAYS"),
            _p("el-002", 1, "The diagram summarizes the pathway just described."),
            Element(
                id="el-003",
                type="image",
                page_number=3,
                ref="images/neural-pathway.png",
                caption="Summary of the neural pathway",
            ),
            _h("el-004", 4, "ENDOCRINE FEEDBACK"),
            _p("el-005", 4, "Hormonal feedback controls endocrine output."),
        ],
    )
    cases.append(
        {
            "id": "image-on-topic-boundary",
            "kind": "ir",
            "description": "A standalone image remains attached to the preceding topic.",
            "source": source,
            "expected": {"topic_boundaries_after": ["el-003"]},
        }
    )

    return cases


def _build_blank_pdf() -> None:
    import pymupdf

    path = SOURCE_DIR / "blank-middle-page.pdf"
    doc = pymupdf.open()
    doc.set_metadata(
        {
            "title": "Golden blank-page fixture",
            "author": "ducsplit",
            "creationDate": "D:20000101000000Z",
            "modDate": "D:20000101000000Z",
        }
    )
    page = doc.new_page()
    page.insert_text((72, 72), "FIRST TOPIC\nThis page contains extractable text.")
    doc.new_page()  # Deliberately blank page 2.
    page = doc.new_page()
    page.insert_text((72, 72), "SECOND TOPIC\nText resumes after the blank page.")
    doc.save(path, garbage=4, deflate=True, no_new_id=True)
    doc.close()


def _normalize_zip(path: Path) -> None:
    """Rewrite an Office zip with stable ordering and timestamps."""
    from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

    temporary = path.with_suffix(path.suffix + ".tmp")
    with ZipFile(path, "r") as source:
        entries = [(info, source.read(info.filename)) for info in source.infolist()]
    with ZipFile(temporary, "w") as target:
        for original, payload in sorted(entries, key=lambda item: item[0].filename):
            info = ZipInfo(original.filename, date_time=(2000, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            target.writestr(info, payload)
    temporary.replace(path)


def _build_docx() -> None:
    from datetime import datetime, timezone

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    import pymupdf

    image_path = SOURCE_DIR / "fixture-image.png"
    pixmap = pymupdf.Pixmap(pymupdf.csRGB, (0, 0, 64, 64), 0)
    pixmap.clear_with(0x4B86B4)
    pixmap.save(image_path)

    document = Document()
    fixed_time = datetime(2000, 1, 1, tzinfo=timezone.utc)
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    document.core_properties.last_modified_by = "ducsplit"
    document.add_heading("DOCX STRUCTURE FIXTURE", level=1)
    document.add_paragraph("First standard Word bullet", style="List Bullet")
    document.add_paragraph("Second standard Word bullet", style="List Bullet")
    document.add_paragraph("A paragraph before the standalone image.")
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(0.5))
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "B"
    output = SOURCE_DIR / "list-and-standalone-image.docx"
    document.save(output)
    _normalize_zip(output)


def main() -> None:
    IR_DIR.mkdir(parents=True, exist_ok=True)
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)

    cases = _build_ir_cases()
    _build_blank_pdf()
    _build_docx()

    cases.extend(
        [
            {
                "id": "pdf-blank-middle-page",
                "kind": "pdf",
                "description": "A blank PDF page is flagged without crashing the parser.",
                "source": "source/blank-middle-page.pdf",
                "expected": {"total_pages": 3, "skipped_pages": [2]},
            },
            {
                "id": "docx-list-and-standalone-image",
                "kind": "docx",
                "description": "Standard Word bullets and an image-only paragraph are preserved.",
                "source": "source/list-and-standalone-image.docx",
                "expected": {
                    "minimum_element_counts": {"list": 1, "image": 1, "table": 1},
                    "list_items": [
                        "First standard Word bullet",
                        "Second standard Word bullet",
                    ],
                },
            },
        ]
    )

    corpus = {
        "schema_version": 1,
        "purpose": "Phase-zero regression and acceptance corpus",
        "desired_page_policy": {
            "target_min_pages": 5,
            "preferred_max_pages": 12,
            "soft_max_pages": 13,
            "hard_max_pages": 20,
            "topic_change_overrides_minimum": True,
            "extension_after_soft_max_requires_semantic_evidence": True,
        },
        "cases": cases,
    }
    (GOLDEN / "corpus.json").write_text(
        json.dumps(corpus, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
